"""
Export Generator Module
Generates KKR-DR01 and KKR-OS01 in Excel and PDF formats
with QR code barcode for reviewer identity
"""
import io
import json
import hashlib
import os
import base64
from datetime import datetime
from modules.data_loader import get_icd_dict
from modules.db_manager import get_audit_db
from modules.rule_engine import validate_case

def get_icd_desc_robust(code, icd_dict):
    if not code: return '-'
    c = str(code).strip()
    while len(c) >= 3:
        if c in icd_dict: return icd_dict[c]
        if c.replace('.', '') in icd_dict: return icd_dict[c.replace('.', '')]
        c = c[:-1]
    return '-'

# --- QR Code ---
try:
    import qrcode
    from qrcode.image.pil import PilImage
    QR_AVAILABLE = True
except ImportError:
    QR_AVAILABLE = False

# --- PDF (ReportLab) ---
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm, mm
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (SimpleDocTemplate, Table, TableStyle, Paragraph,
                                     Spacer, HRFlowable, Image as RLImage, KeepTogether)
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# --- Excel (openpyxl) ---
try:
    import openpyxl
    from openpyxl.styles import (Font, Fill, PatternFill, Alignment, Border, Side,
                                  GradientFill)
    from openpyxl.utils import get_column_letter
    from openpyxl.drawing.image import Image as XLImage
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

# --- PIL for QR + image ---
try:
    from PIL import Image as PILImage
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False


# ============================================================
# QR Code Generation
# ============================================================

def generate_qr_payload(kkr_data, doc_type="KKR-DR01"):
    """Generate payload string for QR code"""
    sep = kkr_data.get('sep', '')
    reviewer = kkr_data.get('reviewer', 'Reviewer')
    rs = kkr_data.get('nama_rs', '')
    ts = kkr_data.get('tanggal_review', datetime.now().isoformat())

    # Hash for integrity
    raw = f"{doc_type}|{sep}|{reviewer}|{ts}"
    hash_val = hashlib.sha256(raw.encode()).hexdigest()[:12]

    payload = {
        "doc": doc_type,
        "sep": sep[:20] if sep else '',
        "reviewer": reviewer[:30] if reviewer else '',
        "rs": rs[:30] if rs else '',
        "timestamp": ts[:19] if ts else '',
        "hash": hash_val
    }
    return json.dumps(payload, ensure_ascii=False)


def generate_qr_image_bytes(kkr_data, doc_type="KKR-DR01", size_px=150):
    """Generate QR code as PNG bytes"""
    if not QR_AVAILABLE or not PIL_AVAILABLE:
        return None

    try:
        payload = generate_qr_payload(kkr_data, doc_type)
        qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.constants.ERROR_CORRECT_M,
            box_size=3,
            border=2,
        )
        qr.add_data(payload)
        qr.make(fit=True)
        img = qr.make_image(fill_color="#003d7a", back_color="white")
        img = img.resize((size_px, size_px), PILImage.LANCZOS if hasattr(PILImage, 'LANCZOS') else PILImage.ANTIALIAS)

        buf = io.BytesIO()
        img.save(buf, format='PNG')
        buf.seek(0)
        return buf.read()
    except Exception as e:
        print(f"[QR] Error: {e}")
        return None


def generate_qr_base64(kkr_data, doc_type="KKR-DR01"):
    """Return QR code as base64 data URI"""
    img_bytes = generate_qr_image_bytes(kkr_data, doc_type)
    if img_bytes:
        b64 = base64.b64encode(img_bytes).decode()
        return f"data:image/png;base64,{b64}"
    return None


# ============================================================
# Excel Export
# ============================================================

def _xl_color(hex_color):
    """Helper to create openpyxl color"""
    return hex_color.lstrip('#')


def _border(style='thin'):
    s = Side(style=style, color='CCCCCC')
    return Border(left=s, right=s, top=s, bottom=s)


def _thick_border():
    thick = Side(style='medium', color='003d7a')
    thin = Side(style='thin', color='CCCCCC')
    return Border(left=thick, right=thick, top=thick, bottom=thick)


def export_kkr_dr01_excel(kkr_data, validate_data=None):
    """Export KKR-DR01 to Excel bytes"""
    if not EXCEL_AVAILABLE:
        raise ImportError("openpyxl not available")

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "KKR-DR01"

    # Page setup
    ws.page_setup.paperSize = ws.PAPERSIZE_A4
    ws.page_setup.orientation = ws.ORIENTATION_PORTRAIT
    ws.page_margins.left = 0.5
    ws.page_margins.right = 0.5
    ws.page_margins.top = 0.7
    ws.page_margins.bottom = 0.7

    # Column widths
    col_widths = [4, 18, 18, 18, 18, 18, 18, 20, 20]
    for i, w in enumerate(col_widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    row = 1

    # === HEADER ===
    # Kemenkes logo area
    ws.merge_cells(f'A{row}:B{row+3}')
    cell = ws[f'A{row}']
    cell.value = "KEMENTERIAN\nKESEHATAN\nREPUBLIK INDONESIA"
    cell.font = Font(name='Calibri', bold=True, size=8, color='FFFFFF')
    cell.fill = PatternFill("solid", fgColor=_xl_color('#003d7a'))
    cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

    ws.merge_cells(f'C{row}:G{row+3}')
    cell = ws[f'C{row}']
    cell.value = "KERTAS KERJA REVIEWER – DESK REVIEW\n(KKR-DR01)\nAUDIT CODING DAN VERIFIKASI DUAL CODING\nTransisi INA-CBG menuju iDRG"
    cell.font = Font(name='Calibri', bold=True, size=12, color='FFFFFF')
    cell.fill = PatternFill("solid", fgColor=_xl_color('#003d7a'))
    cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

    # Meta info
    ws.merge_cells(f'H{row}:I{row}')
    ws[f'H{row}'].value = f"KODE DOKUMEN : KKR-DR01"
    ws[f'H{row}'].font = Font(name='Calibri', size=8, bold=True)
    ws[f'H{row}'].alignment = Alignment(horizontal='left', vertical='center')

    ws.merge_cells(f'H{row+1}:I{row+1}')
    ws[f'H{row+1}'].value = f"VERSI : 1.0"
    ws[f'H{row+1}'].font = Font(name='Calibri', size=8)

    ws.merge_cells(f'H{row+2}:I{row+2}')
    ws[f'H{row+2}'].value = f"TANGGAL : {datetime.now().strftime('%d/%m/%Y')}"
    ws[f'H{row+2}'].font = Font(name='Calibri', size=8)

    ws.merge_cells(f'H{row+3}:I{row+3}')
    ws[f'H{row+3}'].value = "HALAMAN : 1 dari 1"
    ws[f'H{row+3}'].font = Font(name='Calibri', size=8)

    # QR Code in header
    qr_bytes = generate_qr_image_bytes(kkr_data, "KKR-DR01", 100)
    if qr_bytes:
        qr_buf = io.BytesIO(qr_bytes)
        img = XLImage(qr_buf)
        img.width = 80
        img.height = 80
        ws.add_image(img, f'H{row}')

    row += 5

    def section_header(ws, row, num, title, col_end='I'):
        ws.merge_cells(f'A{row}:{col_end}{row}')
        cell = ws[f'A{row}']
        cell.value = f"  {num}. {title}"
        cell.font = Font(name='Calibri', bold=True, size=9, color='FFFFFF')
        cell.fill = PatternFill("solid", fgColor=_xl_color('#1e3a5f'))
        cell.alignment = Alignment(horizontal='left', vertical='center')
        ws.row_dimensions[row].height = 18
        return row + 1

    def data_row(ws, row, label, value, col_end='I'):
        ws.merge_cells(f'A{row}:B{row}')
        ws[f'A{row}'].value = label
        ws[f'A{row}'].font = Font(name='Calibri', size=9, color='666666')
        ws[f'A{row}'].alignment = Alignment(horizontal='left', vertical='center')

        ws.merge_cells(f'C{row}:{col_end}{row}')
        ws[f'C{row}'].value = str(value) if value else '—'
        ws[f'C{row}'].font = Font(name='Calibri', size=9, bold=True)
        ws[f'C{row}'].alignment = Alignment(horizontal='left', vertical='center')

        for col_letter in ['A', 'B', 'C']:
            ws[f'{col_letter}{row}'].border = _border()
        ws.row_dimensions[row].height = 16
        return row + 1

    # === SECTION 1: Identitas ===
    case = kkr_data.get('case', kkr_data)
    row = section_header(ws, row, 1, "IDENTITAS KLAIM (DATA KLAIM DATA CENTER)")
    row = data_row(ws, row, "Nomor SEP", kkr_data.get('sep', case.get('sep', '—')))
    row = data_row(ws, row, "Nomor Klaim", '—')
    row = data_row(ws, row, "Fasilitas Kesehatan (FPKTL)", kkr_data.get('nama_rs', case.get('nama_rs', '—')))
    row = data_row(ws, row, "Kode FPKTL", kkr_data.get('kode_rs', case.get('kode_rs', '—')))
    row = data_row(ws, row, "Tanggal Pelayanan", case.get('discharge_date', '—'))
    row = data_row(ws, row, "Kelas Rawat", f"Kelas {case.get('kelas_rawat', '—')}")
    row = data_row(ws, row, "Length of Stay (LOS)", f"{case.get('alos', '—')} hari")

    row += 1

    # === SECTION 2: Grouping & Tarif ===
    row = section_header(ws, row, 2, "INFORMASI GROUPING DAN TARIF")

    # INA-CBG vs iDRG header
    ws.merge_cells(f'A{row}:B{row}')
    ws[f'A{row}'].value = "Komponen"
    ws[f'A{row}'].font = Font(name='Calibri', bold=True, size=9, color='FFFFFF')
    ws[f'A{row}'].fill = PatternFill("solid", fgColor=_xl_color('#2563EB'))
    ws[f'A{row}'].alignment = Alignment(horizontal='center', vertical='center')

    ws.merge_cells(f'C{row}:E{row}')
    ws[f'C{row}'].value = "INA-CBG"
    ws[f'C{row}'].font = Font(name='Calibri', bold=True, size=9, color='FFFFFF')
    ws[f'C{row}'].fill = PatternFill("solid", fgColor=_xl_color('#0891b2'))
    ws[f'C{row}'].alignment = Alignment(horizontal='center', vertical='center')

    ws.merge_cells(f'F{row}:I{row}')
    ws[f'F{row}'].value = "iDRG"
    ws[f'F{row}'].font = Font(name='Calibri', bold=True, size=9, color='FFFFFF')
    ws[f'F{row}'].fill = PatternFill("solid", fgColor=_xl_color('#7c3aed'))
    ws[f'F{row}'].alignment = Alignment(horizontal='center', vertical='center')
    row += 1

    # Grouping rows
    for label, ina_val, idrg_val in [
        ("Kode", case.get('inacbg', '—'), case.get('idrg_code', '—')),
        ("Deskripsi", (case.get('deskripsi_inacbg', '') or '')[:50], (case.get('deskripsi_idrg', '') or '')[:50]),
    ]:
        ws.merge_cells(f'A{row}:B{row}')
        ws[f'A{row}'].value = label
        ws[f'A{row}'].font = Font(name='Calibri', size=9, color='555555')
        ws[f'A{row}'].alignment = Alignment(horizontal='left', vertical='center')

        ws.merge_cells(f'C{row}:E{row}')
        ws[f'C{row}'].value = str(ina_val) if ina_val else '—'
        ws[f'C{row}'].font = Font(name='Calibri', size=9, bold=True, color='0369A1')
        ws[f'C{row}'].alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)

        ws.merge_cells(f'F{row}:I{row}')
        ws[f'F{row}'].value = str(idrg_val) if idrg_val else '—'
        ws[f'F{row}'].font = Font(name='Calibri', size=9, bold=True, color='6D28D9')
        ws[f'F{row}'].alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)

        for c in ['A', 'C', 'F']:
            ws[f'{c}{row}'].border = _border()
        ws.row_dimensions[row].height = 16
        row += 1

    # Tarif
    tarif_ina = case.get('tarif_inacbg', 0) or 0
    tarif_rs = case.get('tarif_rs', 0) or 0
    selisih = tarif_ina - tarif_rs

    for label, val, color in [
        ("Tarif Klaim (INA-CBG)", f"Rp {tarif_ina:,.0f}", 'FF0EA5E9'),
        ("Tarif RS (Tarif Standar)", f"Rp {tarif_rs:,.0f}", 'FF7C3AED'),
        ("Selisih", f"Rp {abs(selisih):,.0f} ({'INA > RS' if selisih >= 0 else 'RS > INA'})", 'FFDC2626' if selisih > 0 else 'FF059669'),
    ]:
        ws.merge_cells(f'A{row}:D{row}')
        ws[f'A{row}'].value = label
        ws[f'A{row}'].font = Font(name='Calibri', size=9, color='555555')
        ws[f'A{row}'].alignment = Alignment(horizontal='left', vertical='center')

        ws.merge_cells(f'E{row}:I{row}')
        ws[f'E{row}'].value = val
        ws[f'E{row}'].font = Font(name='Calibri', size=10, bold=True, color=color[2:])
        ws[f'E{row}'].alignment = Alignment(horizontal='right', vertical='center')
        for c in ['A', 'E']:
            ws[f'{c}{row}'].border = _border()
        ws.row_dimensions[row].height = 16
        row += 1

    row += 1

    # === SECTION 3: Diagnosis & Prosedur ===
    row = section_header(ws, row, 3, "INPUT DATA KLAIM – DIAGNOSA & PROSEDUR")

    # Diagnosis header
    diag_header = ['No.', 'Kode Diagnosa (INA-CBG)', 'Kode Diagnosa (iDRG)', 'No.', 'Kode Prosedur (INA-CBG)', 'Kode Prosedur (iDRG)']
    for ci, h in enumerate(diag_header, 1):
        cell = ws.cell(row, ci, h)
        cell.font = Font(name='Calibri', bold=True, size=8, color='FFFFFF')
        cell.fill = PatternFill("solid", fgColor=_xl_color('#1e40af'))
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = _border()
    ws.row_dimensions[row].height = 16
    row += 1

    # Diagnosis rows
    diag_codes = [c.strip() for c in str(case.get('diaglist', '') or '').split(';') if c.strip()]
    proc_codes = [c.strip() for c in str(case.get('proclist', '') or '').split(';') if c.strip()]
    max_rows = max(10, len(diag_codes), len(proc_codes))

    triggered_codes = set()
    if validate_data and validate_data.get('triggered_rules'):
        for r in validate_data['triggered_rules']:
            evidence = r.get('evidence', '')
            import re
            matches = re.findall(r'[A-Z]\d{2}(\.\d+)?', evidence)
            triggered_codes.update(matches)

    icd_dict = get_icd_dict()

    for i in range(max_rows):
        d_code = diag_codes[i] if i < len(diag_codes) else ''
        p_code = proc_codes[i] if i < len(proc_codes) else ''
        is_triggered = d_code in triggered_codes
        
        d_desc = f"{d_code} - {get_icd_desc_robust(d_code, icd_dict)}" if d_code else ''
        p_desc = f"{p_code} - {get_icd_desc_robust(p_code, icd_dict)}" if p_code else ''

        row_data = [i+1, d_desc, d_desc, i+1, p_desc, p_desc]
        for ci, val in enumerate(row_data, 1):
            cell = ws.cell(row, ci, val)
            cell.font = Font(name='Calibri', size=8,
                           bold=(ci == 2 and bool(d_code)),
                           color='DC2626' if (ci == 2 and is_triggered) else '0369A1' if ci in [2] else '6D28D9' if ci in [3, 5, 6] else '333333')
            cell.alignment = Alignment(horizontal='center' if ci in [1, 4] else 'left', vertical='center', wrap_text=True)
            cell.border = _border()
        ws.row_dimensions[row].height = 24 if (d_code or p_code) else 15
        row += 1

    row += 1

    # === SECTION 4: Ringkasan Temuan ===
    row = section_header(ws, row, 4, "RINGKASAN TEMUAN (HASIL VALIDASI OTOMATIS BERDASARKAN RULE)")

    rules = (validate_data or {}).get('triggered_rules', kkr_data.get('triggered_rules', []))

    if not rules:
        ws.merge_cells(f'A{row}:I{row}')
        ws[f'A{row}'].value = "✅ Tidak ada temuan – Semua aturan validasi terpenuhi"
        ws[f'A{row}'].font = Font(name='Calibri', size=9, color='059669', bold=True)
        ws[f'A{row}'].alignment = Alignment(horizontal='center', vertical='center')
        ws.row_dimensions[row].height = 16
        row += 1
    else:
        # Rules table header
        rule_headers = ['No.', 'Rule ID', 'Nama Aturan', 'Hasil Validasi', 'Severity', 'Evidence', 'Rekomendasi']
        rule_col_widths_map = {1: 4, 2: 14, 3: 22, 4: 14, 5: 10, 6: 22, 7: 24}
        for ci, (h, w) in enumerate(zip(rule_headers, rule_col_widths_map.values()), 1):
            cell = ws.cell(row, ci, h)
            cell.font = Font(name='Calibri', bold=True, size=8, color='FFFFFF')
            cell.fill = PatternFill("solid", fgColor=_xl_color('#dc2626'))
            cell.alignment = Alignment(horizontal='center', vertical='center')
            cell.border = _border()
        ws.row_dimensions[row].height = 16
        row += 1

        for i, rule in enumerate(rules):
            sev = rule.get('severity', '')
            sev_color = 'DC2626' if sev == 'High' else 'D97706' if sev == 'Medium' else '059669'
            row_data = [
                i+1, rule.get('rule_id', ''), rule.get('nama_aturan', ''),
                'Terindikasi', sev, rule.get('evidence', ''), rule.get('rekomendasi_reviewer', '')
            ]
            for ci, val in enumerate(row_data, 1):
                cell = ws.cell(row, ci, val)
                cell.font = Font(name='Calibri', size=8,
                               color=sev_color if ci == 5 else '333333',
                               bold=(ci == 5))
                cell.alignment = Alignment(horizontal='center' if ci in [1, 4, 5] else 'left',
                                         vertical='center', wrap_text=True)
                cell.border = _border()
            ws.row_dimensions[row].height = 30
            row += 1

    row += 1

    # === SECTION 5 & 6: Analisis & Keputusan ===
    row = section_header(ws, row, 5, "ANALISIS REVIEWER")
    ws.merge_cells(f'A{row}:I{row+2}')
    ws[f'A{row}'].value = kkr_data.get('analisis_reviewer', '')
    ws[f'A{row}'].font = Font(name='Calibri', size=9)
    ws[f'A{row}'].alignment = Alignment(horizontal='left', vertical='top', wrap_text=True)
    ws[f'A{row}'].border = _border()
    ws.row_dimensions[row].height = 60
    row += 3

    row = section_header(ws, row, 6, "KEPUTUSAN REVIEWER")

    keputusan = kkr_data.get('keputusan_reviewer', '')
    tingkat = kkr_data.get('tingkat_keyakinan', '')
    alasan = kkr_data.get('alasan', '')

    for label, val in [
        ("Keputusan", keputusan),
        ("Tingkat Keyakinan", tingkat),
        ("Alasan", alasan),
    ]:
        row = data_row(ws, row, label, val)

    row += 1

    # === SECTION 7: Paraf ===
    row = section_header(ws, row, 7, "PARAF REVIEWER")

    ws[f'A{row}'].value = "Reviewer,"
    ws[f'D{row}'].value = "Ketua Tim Reviewer,"
    for cell in [ws[f'A{row}'], ws[f'D{row}']]:
        cell.font = Font(name='Calibri', size=9, bold=True)

    row += 3

    ws[f'A{row}'].value = f"( {kkr_data.get('reviewer', '________________')} )"
    ws[f'D{row}'].value = f"( {kkr_data.get('ketua_tim', '________________')} )"
    for cell in [ws[f'A{row}'], ws[f'D{row}']]:
        cell.font = Font(name='Calibri', size=9)
        cell.alignment = Alignment(horizontal='center')

    row += 1
    ws[f'A{row}'].value = f"Tanggal: {kkr_data.get('tgl_reviewer', '___/___/______')}"
    ws[f'D{row}'].value = f"Tanggal: {kkr_data.get('tgl_ketua', '___/___/______')}"

    # QR Code at bottom
    qr_bytes = generate_qr_image_bytes(kkr_data, "KKR-DR01", 120)
    if qr_bytes:
        qr_buf = io.BytesIO(qr_bytes)
        qr_img = XLImage(qr_buf)
        qr_img.width = 100
        qr_img.height = 100
        ws.add_image(qr_img, f'H{row-3}')

    # Save
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()


# ============================================================
# PDF Export (ReportLab)
# ============================================================

def export_kkr_dr01_pdf(kkr_data, validate_data=None):
    """Export KKR-DR01 to PDF bytes"""
    if not PDF_AVAILABLE:
        raise ImportError("reportlab not available")

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        rightMargin=1.5*cm,
        leftMargin=1.5*cm,
        topMargin=1.5*cm,
        bottomMargin=1.5*cm
    )

    styles = getSampleStyleSheet()

    # Custom styles
    h_style = ParagraphStyle('KKRHeading', fontName='Helvetica-Bold', fontSize=14,
                              textColor=colors.white, alignment=TA_CENTER,
                              backColor=colors.HexColor('#003d7a'), spaceAfter=2)
    sub_style = ParagraphStyle('KKRSub', fontName='Helvetica-Bold', fontSize=10,
                               textColor=colors.HexColor('#06b6d4'), alignment=TA_CENTER,
                               spaceAfter=2)
    section_style = ParagraphStyle('Section', fontName='Helvetica-Bold', fontSize=9,
                                   textColor=colors.white, backColor=colors.HexColor('#1e3a5f'),
                                   spaceAfter=3, spaceBefore=6, leftIndent=4)
    normal_style = ParagraphStyle('Normal9', fontName='Helvetica', fontSize=8,
                                  textColor=colors.HexColor('#333333'), spaceAfter=2)
    bold_style = ParagraphStyle('Bold9', fontName='Helvetica-Bold', fontSize=8,
                                textColor=colors.HexColor('#111111'), spaceAfter=2)
    muted_style = ParagraphStyle('Muted', fontName='Helvetica', fontSize=7,
                                 textColor=colors.HexColor('#888888'))

    story = []

    # --- Header ---
    case = kkr_data.get('case', kkr_data)
    sep = kkr_data.get('sep', case.get('sep', '—'))
    nama_rs = kkr_data.get('nama_rs', case.get('nama_rs', '—'))

    # Header table with QR code
    qr_bytes = generate_qr_image_bytes(kkr_data, "KKR-DR01", 80)
    qr_img_el = None
    if qr_bytes:
        qr_io = io.BytesIO(qr_bytes)
        qr_img_el = RLImage(qr_io, width=2.2*cm, height=2.2*cm)

    header_data = [[
        Paragraph("KEMENTERIAN<br/>KESEHATAN<br/>REPUBLIK INDONESIA",
                  ParagraphStyle('Logo', fontName='Helvetica-Bold', fontSize=7,
                                 textColor=colors.white, alignment=TA_CENTER)),
        Paragraph("KERTAS KERJA REVIEWER – DESK REVIEW<br/>(KKR-DR01)<br/>AUDIT CODING DAN VERIFIKASI DUAL CODING<br/>Transisi INA-CBG menuju Indonesian Diagnosis Related Groups (iDRG)",
                  ParagraphStyle('Title', fontName='Helvetica-Bold', fontSize=10,
                                 textColor=colors.white, alignment=TA_CENTER)),
        Paragraph(f"KODE DOKUMEN : KKR-DR01<br/>VERSI : 1.0<br/>TGL BERLAKU : {datetime.now().strftime('%d/%m/%Y')}<br/>HALAMAN : 1 dari 1",
                  ParagraphStyle('Meta', fontName='Helvetica', fontSize=7,
                                 textColor=colors.white, alignment=TA_LEFT)),
        qr_img_el or Paragraph("", normal_style)
    ]]
    header_table = Table(header_data, colWidths=[3*cm, 9*cm, 4.5*cm, 2.5*cm])
    header_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#003d7a')),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('LEFTPADDING', (0, 0), (-1, -1), 6),
        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
        ('LINEBELOW', (0, 0), (-1, -1), 2, colors.HexColor('#06b6d4')),
    ]))
    story.append(header_table)
    story.append(Spacer(1, 8))

    def section_hdr(text):
        return Paragraph(f"  {text}", section_style)

    def info_row(label, value, color=None):
        return [
            Paragraph(label, muted_style),
            Paragraph(str(value or '—'), bold_style if not color else
                      ParagraphStyle('Val', fontName='Helvetica-Bold', fontSize=8, textColor=color))
        ]

    # === Section 1: Identitas ===
    story.append(section_hdr("1.  IDENTITAS KLAIM (DATA KLAIM DATA CENTER)"))

    id_data = [
        info_row("Nomor SEP", sep),
        info_row("Nomor Klaim", '—'),
        info_row("Fasilitas Kesehatan (FPKTL)", nama_rs),
        info_row("Kode FPKTL", kkr_data.get('kode_rs', case.get('kode_rs', '—'))),
        info_row("Tanggal Pelayanan", case.get('discharge_date', '—')),
        info_row("Kelas Rawat", f"Kelas {case.get('kelas_rawat', '—')}"),
        info_row("Length of Stay (LOS)", f"{case.get('alos', '—')} hari"),
    ]

    id_table = Table(id_data, colWidths=[5.5*cm, 13.5*cm])
    id_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#666666')),
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#f8fafc')),
        ('ROWBACKGROUNDS', (0, 0), (-1, -1), [colors.white, colors.HexColor('#f0f4f8')]),
        ('GRID', (0, 0), (-1, -1), 0.3, colors.HexColor('#e0e8f0')),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ('LEFTPADDING', (0, 0), (-1, -1), 6),
    ]))
    story.append(id_table)
    story.append(Spacer(1, 6))

    # === Section 2: Grouping & Tarif ===
    story.append(section_hdr("2.  INFORMASI GROUPING DAN TARIF"))

    tarif_ina = case.get('tarif_inacbg', 0) or 0
    tarif_rs = case.get('tarif_rs', 0) or 0
    selisih = tarif_ina - tarif_rs

    grouping_data = [
        [Paragraph("Komponen", bold_style), Paragraph("INA-CBG", bold_style), Paragraph("iDRG", bold_style)],
        [Paragraph("Kode", muted_style), Paragraph(case.get('inacbg', '—'), bold_style), Paragraph(case.get('idrg_code', '—'), bold_style)],
        [Paragraph("Deskripsi", muted_style),
         Paragraph((case.get('deskripsi_inacbg', '') or '')[:60], normal_style),
         Paragraph((case.get('deskripsi_idrg', '') or '')[:60], normal_style)],
        [Paragraph("Tarif Klaim (INA-CBG)", muted_style), Paragraph(f"Rp {tarif_ina:,.0f}", bold_style), Paragraph('', normal_style)],
        [Paragraph("Tarif RS", muted_style), Paragraph(f"Rp {tarif_rs:,.0f}", bold_style), Paragraph('', normal_style)],
        [Paragraph("Selisih", ParagraphStyle('SelisihLbl', fontName='Helvetica-Bold', fontSize=8, textColor=colors.HexColor('#555555'))),
         Paragraph(f"Rp {abs(selisih):,.0f} ({'INA > RS' if selisih >= 0 else 'RS > INA'})",
                   ParagraphStyle('SelisihVal', fontName='Helvetica-Bold', fontSize=9,
                                  textColor=colors.HexColor('#DC2626') if selisih > 0 else colors.HexColor('#059669'))),
         Paragraph('', normal_style)],
    ]
    g_table = Table(grouping_data, colWidths=[5.5*cm, 8.5*cm, 5*cm])
    g_table.setStyle(TableStyle([
        ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1e3a5f')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f0f4f8')]),
        ('GRID', (0, 0), (-1, -1), 0.3, colors.HexColor('#e0e8f0')),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ('LEFTPADDING', (0, 0), (-1, -1), 6),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    story.append(g_table)
    story.append(Spacer(1, 6))

    # === Section 3: Diagnosa & Prosedur ===
    story.append(section_hdr("3.  INPUT DATA KLAIM (BERDASARKAN DATA KLAIM DATA CENTER)"))

    diag_codes = [c.strip() for c in str(case.get('diaglist', '') or '').split(';') if c.strip()]
    proc_codes = [c.strip() for c in str(case.get('proclist', '') or '').split(';') if c.strip()]

    triggered_rules = (validate_data or {}).get('triggered_rules', kkr_data.get('triggered_rules', []))
    triggered_set = set()
    for r in triggered_rules:
        import re
        matches = re.findall(r'[A-Z]\d{2}(?:\.\d+)?', r.get('evidence', ''))
        triggered_set.update(matches)

    max_rows = max(10, max(len(diag_codes), len(proc_codes)))
    icd_dict = get_icd_dict()

    # Table Header
    dp_rows = [
        [Paragraph("No.", bold_style), Paragraph("DIAGNOSA", bold_style), "", Paragraph("PROSEDUR", bold_style), ""],
        ["", Paragraph("INA-CBG / iDRG<br/>Kode ICD", bold_style), Paragraph("Deskripsi", bold_style), Paragraph("INA-CBG / iDRG<br/>Kode ICD-9-CM", bold_style), Paragraph("Deskripsi", bold_style)]
    ]

    for i in range(max_rows):
        d_code = diag_codes[i] if i < len(diag_codes) else ''
        p_code = proc_codes[i] if i < len(proc_codes) else ''
        
        d_trig = d_code in triggered_set
        p_trig = p_code in triggered_set
        
        d_style = ParagraphStyle('CT', fontName='Helvetica-Bold', fontSize=8, textColor=colors.HexColor('#DC2626')) if d_trig else bold_style
        p_style = ParagraphStyle('CT', fontName='Helvetica-Bold', fontSize=8, textColor=colors.HexColor('#DC2626')) if p_trig else bold_style
        
        d_desc = get_icd_desc_robust(d_code, icd_dict) if d_code else '-'
        p_desc = get_icd_desc_robust(p_code, icd_dict) if p_code else '-'
        
        dp_rows.append([
            Paragraph(str(i+1), normal_style),
            Paragraph(d_code or '-', d_style),
            Paragraph(d_desc, ParagraphStyle('Desc', fontName='Helvetica', fontSize=7, leading=8)),
            Paragraph(p_code or '-', p_style),
            Paragraph(p_desc, ParagraphStyle('Desc', fontName='Helvetica', fontSize=7, leading=8))
        ])

    dp_table = Table(dp_rows, colWidths=[1*cm, 3.2*cm, 5.3*cm, 3.2*cm, 5.3*cm])
    dp_table.setStyle(TableStyle([
        ('SPAN', (0, 0), (0, 1)), # Span No.
        ('SPAN', (1, 0), (2, 0)), # Span DIAGNOSA
        ('SPAN', (3, 0), (4, 0)), # Span PROSEDUR
        ('BACKGROUND', (0, 0), (-1, 1), colors.HexColor('#0369A1')),
        ('TEXTCOLOR', (0, 0), (-1, 1), colors.white),
        ('ALIGN', (0, 0), (-1, 1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('GRID', (0, 0), (-1, -1), 0.3, colors.HexColor('#BAE6FD')),
        ('ROWBACKGROUNDS', (2, 2), (-1, -1), [colors.white, colors.HexColor('#f0f9ff')]),
        ('FONTSIZE', (0, 0), (-1, -1), 7.5),
        ('TOPPADDING', (0, 0), (-1, -1), 3),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ('ALIGN', (0, 2), (0, -1), 'CENTER'),
        ('ALIGN', (1, 2), (1, -1), 'CENTER'),
        ('ALIGN', (3, 2), (3, -1), 'CENTER'),
    ]))
    
    story.append(dp_table)
    story.append(Spacer(1, 4))
    story.append(Paragraph("Catatan: Isi sesuai urutan yang tercantum pada data klaim.", ParagraphStyle('Note', fontName='Helvetica-Oblique', fontSize=7, textColor=colors.gray)))
    story.append(Spacer(1, 10))

    # === Section 4: Ringkasan Temuan ===
    story.append(section_hdr("4.  RINGKASAN TEMUAN (HASIL VALIDASI OTOMATIS BERDASARKAN RULE)"))

    if not triggered_rules:
        story.append(Paragraph("✅  Tidak ada temuan – Semua aturan validasi terpenuhi",
                                ParagraphStyle('OK', fontName='Helvetica-Bold', fontSize=9,
                                               textColor=colors.HexColor('#059669'), alignment=TA_CENTER)))
    else:
        rule_data = [[
            Paragraph("No.", bold_style),
            Paragraph("Rule ID", bold_style),
            Paragraph("Nama Aturan", bold_style),
            Paragraph("Hasil", bold_style),
            Paragraph("Severity", bold_style),
            Paragraph("Evidence / Pesan", bold_style),
        ]]
        for i, rule in enumerate(triggered_rules):
            sev = rule.get('severity', '')
            sev_color = colors.HexColor('#DC2626') if sev == 'High' else colors.HexColor('#D97706') if sev == 'Medium' else colors.HexColor('#059669')
            rule_data.append([
                Paragraph(str(i+1), normal_style),
                Paragraph(rule.get('rule_id', ''), ParagraphStyle('RID', fontName='Helvetica-Bold', fontSize=7, textColor=colors.HexColor('#0369A1'))),
                Paragraph(rule.get('nama_aturan', ''), ParagraphStyle('RN', fontName='Helvetica', fontSize=7.5)),
                Paragraph("Terindikasi", ParagraphStyle('TI', fontName='Helvetica-Bold', fontSize=7, textColor=colors.HexColor('#D97706'))),
                Paragraph(sev, ParagraphStyle('SEV', fontName='Helvetica-Bold', fontSize=7, textColor=sev_color)),
                Paragraph((rule.get('evidence', '') + '<br/>' + rule.get('pesan_validasi', ''))[:120],
                          ParagraphStyle('Pesan', fontName='Helvetica', fontSize=7)),
            ])

        rule_table = Table(rule_data, colWidths=[0.7*cm, 2.5*cm, 4*cm, 2*cm, 1.8*cm, 8*cm])
        rule_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#DC2626')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#fef2f2')]),
            ('GRID', (0, 0), (-1, -1), 0.3, colors.HexColor('#FECACA')),
            ('FONTSIZE', (0, 0), (-1, -1), 7.5),
            ('TOPPADDING', (0, 0), (-1, -1), 3), ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ('LEFTPADDING', (0, 0), (-1, -1), 4), ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('ALIGN', (0, 0), (0, -1), 'CENTER'),
        ]))
        story.append(rule_table)

    story.append(Spacer(1, 6))

    # === Section 5 & 6: Analisis & Keputusan ===
    story.append(section_hdr("5.  ANALISIS & KEPUTUSAN REVIEWER"))

    analisis = kkr_data.get('analisis_reviewer', '—')
    keputusan = kkr_data.get('keputusan_reviewer') or kkr_data.get('keputusan', '—')
    tingkat = kkr_data.get('tingkat_keyakinan', '—')
    alasan = kkr_data.get('alasan') or kkr_data.get('alasan_keputusan', '—')

    ak_data = [
        [Paragraph("Analisis Reviewer", muted_style), Paragraph(analisis or '—', normal_style)],
        [Paragraph("Keputusan Reviewer", muted_style), Paragraph(keputusan or '—', bold_style)],
        [Paragraph("Tingkat Keyakinan", muted_style), Paragraph(tingkat or '—', normal_style)],
        [Paragraph("Alasan / Catatan", muted_style), Paragraph(alasan or '—', normal_style)],
    ]
    ak_table = Table(ak_data, colWidths=[4.5*cm, 14.5*cm])
    ak_table.setStyle(TableStyle([
        ('GRID', (0, 0), (-1, -1), 0.3, colors.HexColor('#e0e0e0')),
        ('ROWBACKGROUNDS', (0, 0), (-1, -1), [colors.HexColor('#eff6ff'), colors.white]),
        ('TOPPADDING', (0, 0), (-1, -1), 5), ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ('LEFTPADDING', (0, 0), (-1, -1), 6),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
    ]))
    story.append(ak_table)
    story.append(Spacer(1, 8))

    # === Section 7: Paraf ===
    story.append(section_hdr("6.  PARAF REVIEWER"))

    reviewer = kkr_data.get('reviewer_name') or kkr_data.get('reviewer') or 'System (Auto)'
    ketua = kkr_data.get('ketua_tim_name') or kkr_data.get('ketua_tim') or '________________'
    tgl_rev = kkr_data.get('tanggal_review') or kkr_data.get('tgl_reviewer') or '___/___/______'
    tgl_ket = kkr_data.get('tanggal_ketua') or kkr_data.get('tgl_ketua') or '___/___/______'

    # QR code at bottom right
    paraf_qr = None
    if qr_bytes:
        qr_io2 = io.BytesIO(qr_bytes)
        paraf_qr = RLImage(qr_io2, width=2*cm, height=2*cm)

    paraf_data = [
        [Paragraph("Reviewer,", muted_style), Paragraph("Ketua Tim Reviewer,", muted_style),
         Paragraph("🔏 QR Kode Integritas", ParagraphStyle('QRLbl', fontName='Helvetica', fontSize=7, textColor=colors.HexColor('#888888'), alignment=TA_CENTER))],
        [Paragraph("<br/><br/>", normal_style), Paragraph("<br/><br/>", normal_style), paraf_qr or Paragraph("", normal_style)],
        [Paragraph(f"( {reviewer} )", normal_style), Paragraph(f"( {ketua} )", normal_style),
         Paragraph("", normal_style)],
        [Paragraph(f"Tanggal: {tgl_rev}", muted_style), Paragraph(f"Tanggal: {tgl_ket}", muted_style),
         Paragraph(generate_qr_payload(kkr_data)[:40] + "...",
                   ParagraphStyle('QRInfo', fontName='Helvetica', fontSize=6, textColor=colors.HexColor('#aaaaaa'), alignment=TA_CENTER))],
    ]
    paraf_table = Table(paraf_data, colWidths=[7*cm, 7*cm, 5*cm])
    paraf_table.setStyle(TableStyle([
        ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('GRID', (0, 0), (1, -1), 0.3, colors.HexColor('#e0e0e0')),
        ('TOPPADDING', (0, 0), (-1, -1), 4), ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
    ]))
    story.append(paraf_table)

    # Footer note
    story.append(Spacer(1, 6))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#CBD5E1')))
    story.append(Spacer(1, 4))
    footer_text = ("Formulir ini digunakan untuk Desk Review berdasarkan data klaim pada Data Center. "
                   "Penilaian ini BUKAN penetapan benar/salah pengodean. "
                   "Verifikasi definitif dilakukan pada On-Site Audit.")
    story.append(Paragraph(footer_text, ParagraphStyle('Footer', fontName='Helvetica', fontSize=6.5,
                                                        textColor=colors.HexColor('#94A3B8'), alignment=TA_CENTER)))

    doc.build(story)
    buf.seek(0)
    return buf.read()


import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_lha_word(kode_rs, rs_name, cases, output_path):
    """
    Generate Laporan Hasil Audit (LHA) in Word format based on the guidelines.
    """
    document = Document()
    
    # Title
    heading = document.add_heading('LAPORAN HASIL AUDIT (LHA) KODING JKN', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    document.add_paragraph(f"Nama Rumah Sakit : {rs_name}")
    document.add_paragraph(f"Kode Rumah Sakit : {kode_rs}")
    document.add_paragraph(f"Jumlah Sampel Audit: {len(cases)} Kasus")
    
    document.add_heading('1. Ringkasan Kasus', level=1)
    
    conn = get_audit_db()
    cursor = conn.cursor()
    cursor.execute("SELECT COUNT(*) as total, SUM(CASE WHEN inacbg LIKE '%-0' THEN 1 ELSE 0 END) as rj, SUM(CASE WHEN inacbg NOT LIKE '%-0' THEN 1 ELSE 0 END) as ri FROM datadb.individual_data WHERE kode_rs = ?", (kode_rs,))
    rs_stats = cursor.fetchone()
    total_rs = rs_stats['total']
    rj_rs = rs_stats['rj'] or 0
    ri_rs = rs_stats['ri'] or 0
    
    table1 = document.add_table(rows=5, cols=2)
    table1.style = 'Table Grid'
    
    t1_cells = table1.columns[0].cells
    t1_cells[0].text = 'Uraian'
    t1_cells[1].text = 'Total Kasus'
    t1_cells[2].text = 'Rawat Jalan'
    t1_cells[3].text = 'Rawat Inap'
    t1_cells[4].text = 'Total Kasus Direview'
    
    v1_cells = table1.columns[1].cells
    v1_cells[0].text = 'Jumlah'
    v1_cells[1].text = str(total_rs)
    v1_cells[2].text = str(rj_rs)
    v1_cells[3].text = str(ri_rs)
    v1_cells[4].text = str(len(cases))
    
    document.add_heading('2. Ringkasan Temuan Audit', level=1)
    
    # Load real rules for these cases
    seps = [c['sep'] for c in cases]
    placeholders = ','.join('?' for _ in seps)
    cursor.execute(f"SELECT sep, diaglist, proclist FROM datadb.individual_data WHERE sep IN ({placeholders})", seps)
    rows = cursor.fetchall()
    conn.close()
    
    full_cases_map = {row['sep']: dict(row) for row in rows}
    
    rule_counts = {
        'combination_code': 0,
        'dagger_asterisk': 0,
        'includes_excludes': 0,
        'underlying_manifestation': 0,
        'procedure_validation': 0,
        'unbundling': 0,
        'medical_evidence': 0,
        'administrative_validation': 0,
        'age_validation': 0,
        'los_validation': 0,
    }
    
    for case in cases:
        sep = case['sep']
        if sep in full_cases_map:
            triggered = validate_case(full_cases_map[sep])
            for rule in triggered:
                kat = rule.get('kategori', 'administrative_validation')
                if kat in rule_counts:
                    rule_counts[kat] += 1
                else:
                    rule_counts['administrative_validation'] += 1

    table2 = document.add_table(rows=11, cols=2)
    table2.style = 'Table Grid'
    t2_data = [
        ['Kelompok Aturan', 'Jumlah Temuan'],
        ['Combination Code', str(rule_counts['combination_code'])],
        ['Dagger & Asterisk', str(rule_counts['dagger_asterisk'])],
        ['Includes dan Excludes', str(rule_counts['includes_excludes'])],
        ['Underlying Cause dan Manifestation', str(rule_counts['underlying_manifestation'])],
        ['Procedure Validation', str(rule_counts['procedure_validation'])],
        ['Unbundling dan Omit Code', str(rule_counts['unbundling'])],
        ['Medical Evidence Validation', str(rule_counts['medical_evidence'])],
        ['Administrative Validation', str(rule_counts['administrative_validation'])],
        ['Age Validation', str(rule_counts['age_validation'])],
        ['Length of Stay Validation', str(rule_counts['los_validation'])]
    ]
    for i, row_data in enumerate(t2_data):
        row_cells = table2.rows[i].cells
        row_cells[0].text = row_data[0]
        row_cells[1].text = row_data[1]
        
    document.add_heading('3. Rincian Sampel Audit', level=1)
    
    # Table 17
    table3 = document.add_table(rows=1, cols=5)
    table3.style = 'Table Grid'
    hdr_cells = table3.rows[0].cells
    hdr_cells[0].text = 'No'
    hdr_cells[1].text = 'Nomor SEP'
    hdr_cells[2].text = 'Keputusan'
    hdr_cells[3].text = 'Selisih Tarif'
    hdr_cells[4].text = 'Rekomendasi'
    
    for idx, case in enumerate(cases):
        row_cells = table3.add_row().cells
        row_cells[0].text = str(idx + 1)
        row_cells[1].text = str(case.get('sep', ''))
        row_cells[2].text = str(case.get('keputusan', ''))
        try:
            selisih = float(case.get('tarif_inacbg') or 0) - float(case.get('tarif_rs') or 0)
        except:
            selisih = 0
        row_cells[3].text = f"Rp {selisih:,.0f}"
        row_cells[4].text = str(case.get('rekomendasi_lanjut', ''))
        
    document.add_heading('4. Pengesahan', level=1)
    table4 = document.add_table(rows=2, cols=3)
    table4.style = 'Table Grid'
    t4_hdr = table4.rows[0].cells
    t4_hdr[0].text = 'Disusun oleh'
    t4_hdr[1].text = 'Direviu oleh'
    t4_hdr[2].text = 'Disahkan oleh'
    
    # Add empty space for signature
    t4_val = table4.rows[1].cells
    t4_val[0].text = '\n\n\n(Tim Audit)'
    t4_val[1].text = '\n\n\n(Ketua Tim)'
    t4_val[2].text = '\n\n\n(Pejabat Berwenang)'
    
    document.save(output_path)
    return output_path
